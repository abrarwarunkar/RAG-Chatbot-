import io
import logging
import psutil
import gc
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        # Optimized chunk settings for memory efficiency
        self.chunk_size = 800  # Smaller chunks to reduce memory
        self.chunk_overlap = 100  # Reduced overlap
        self._log_memory_usage("DocumentProcessor initialized")
    
    def _log_memory_usage(self, context: str):
        """Log current memory usage"""
        try:
            process = psutil.Process()
            memory_mb = process.memory_info().rss / 1024 / 1024
            logger.info(f"{context} - Memory usage: {memory_mb:.2f} MB")
        except Exception as e:
            logger.warning(f"Could not log memory usage: {e}")
    
    async def process_document(self, content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Process document with memory optimization"""
        try:
            self._log_memory_usage(f"Processing {filename}")
            
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = await self._extract_pdf_text(content)
            elif filename.lower().endswith('.docx'):
                text = await self._extract_docx_text(content)
            elif filename.lower().endswith('.txt'):
                text = content.decode('utf-8', errors='ignore')
            else:
                raise ValueError(f"Unsupported file type: {filename}")
            
            # Limit document size to prevent memory issues
            max_doc_size = 50000  # 50KB text limit
            if len(text) > max_doc_size:
                logger.warning(f"Document {filename} truncated from {len(text)} to {max_doc_size} characters")
                text = text[:max_doc_size] + "\n[Document truncated due to size limit]"
            
            # Create chunks with memory optimization
            chunks = self._create_chunks(text, filename)
            
            # Cleanup
            del text
            gc.collect()
            
            self._log_memory_usage(f"Processed {filename} into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing {filename}: {str(e)}")
            # Cleanup on error
            gc.collect()
            raise
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF with memory optimization"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            # Limit number of pages to prevent memory issues
            max_pages = 50
            num_pages = min(len(reader.pages), max_pages)
            
            if len(reader.pages) > max_pages:
                logger.warning(f"PDF truncated from {len(reader.pages)} to {max_pages} pages")
            
            text_parts = []
            for page_num in range(num_pages):
                try:
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text_parts.append(page_text)
                    
                    # Cleanup page object
                    del page, page_text
                    
                    # Force cleanup every 10 pages
                    if page_num % 10 == 0:
                        gc.collect()
                        
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num}: {e}")
                    continue
            
            text = "\n\n".join(text_parts)
            
            # Cleanup
            del pdf_file, reader, text_parts
            gc.collect()
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX with memory optimization"""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            # Extract text with limits
            text_parts = []
            max_paragraphs = 200  # Limit paragraphs
            
            for i, paragraph in enumerate(doc.paragraphs[:max_paragraphs]):
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                
                # Force cleanup every 50 paragraphs
                if i % 50 == 0:
                    gc.collect()
            
            if len(doc.paragraphs) > max_paragraphs:
                logger.warning(f"DOCX truncated from {len(doc.paragraphs)} to {max_paragraphs} paragraphs")
            
            text = "\n\n".join(text_parts)
            
            # Cleanup
            del docx_file, doc, text_parts
            gc.collect()
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _create_chunks(self, text: str, filename: str) -> List[Dict[str, Any]]:
        """Create text chunks with memory optimization"""
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            if end < len(text):
                # Find the last sentence boundary within the chunk
                boundary = text.rfind('.', start, end)
                if boundary == -1 or boundary < start + 100:  # Minimum chunk size
                    boundary = text.rfind(' ', start, end)
                if boundary != -1 and boundary > start:
                    end = boundary + 1
            
            # Extract chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Only add non-empty chunks
                chunks.append({
                    "content": chunk_text,
                    "metadata": {
                        "filename": filename,
                        "chunk_index": chunk_index,
                        "start_pos": start,
                        "end_pos": end
                    }
                })
                chunk_index += 1
            
            # Move to next chunk with overlap
            if end >= len(text):
                break
            
            start = max(start + 1, end - self.chunk_overlap)
            
            # Force cleanup every 50 chunks
            if len(chunks) % 50 == 0:
                gc.collect()
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        return chunks